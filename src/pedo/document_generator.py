from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
import subprocess
import os
import re
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
from docx.shared import RGBColor
from docx.enum.section import WD_SECTION

def create_formatted_case_study(input_text, logo_url, output_file, colors=None):
    """
    Creates a well-formatted Word document from structured input text and converts it to PDF.
    
    Args:
        input_text (str): The formatted text content
        logo_url (str): URL to the company logo
        output_file (str): Path to save the PDF
        colors (dict): Dictionary with color settings (e.g. {'border': '000080', 'header': '4472C4'})
    """
    if not output_file.endswith(".pdf"):
        raise ValueError("Output file must have a .pdf extension")

    # Default colors if none provided
    if colors is None:
        colors = {
            'border': '000000',      # Black
            'header': '4472C4',      # Blue
            'footer': '4472C4',      # Blue (same as header by default)
            'heading1': '2F5597',    # Dark blue
            'heading2': '5B9BD5',    # Medium blue
            'accent': '70AD47'       # Green
        }
    
    # Ensure footer color exists (use header color as fallback)
    if 'footer' not in colors and 'header' in colors:
        colors['footer'] = colors['header']
    
    docx_file = output_file.replace(".pdf", ".docx")
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Download and prepare logo
    logo_path = None
    try:
        if logo_url:
            print(f"📥 Downloading logo from URL: {logo_url}")
            response = requests.get(logo_url, timeout=10)
            if response.status_code == 200:
                print(f"✅ Successfully downloaded logo, content length: {len(response.content)} bytes")
                
                # Save the logo temporarily
                logo_path = "temp_logo.png"
                with open(logo_path, "wb") as f:
                    f.write(response.content)
                
                # Verify the image was properly saved
                if os.path.exists(logo_path) and os.path.getsize(logo_path) > 0:
                    print(f"✅ Saved logo to {logo_path}, size: {os.path.getsize(logo_path)} bytes")
                else:
                    print("⚠️ Failed to save logo image properly")
                    logo_path = None
            else:
                print(f"⚠️ Failed to download logo, status code: {response.status_code}")
        else:
            print("⚠️ No logo URL provided, skipping logo addition")
    except Exception as e:
        print(f"⚠️ Failed to add logo: {str(e)}")
    
    # Extract company name from filename if available
    company_name = ""
    if output_file:
        base_name = os.path.basename(output_file)
        if base_name.startswith("case_study_"):
            company_parts = base_name.replace("case_study_", "").split("_")
            if len(company_parts) > 0:
                company_name = " ".join([part.capitalize() for part in company_parts[0].split("_")])
    
    # Process text content with markdown-like formatting
    data_for_graphs = []
    content_block_started = False
    
    for line in input_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        
        # Headings - start on a new page
        if line.startswith("# ") or line.startswith("## ") or line.startswith("### "):
            # If content block has already started, insert a page break
            if content_block_started:
                doc.add_page_break()
            
            content_block_started = True
            
            if line.startswith("### "): # H3
                p = doc.add_paragraph(line[4:], style="Heading 3")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(colors['heading2'])
                
            elif line.startswith("## "): # H2
                p = doc.add_paragraph(line[3:], style="Heading 2")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(colors['heading2'])
                
            elif line.startswith("# "): # H1
                p = doc.add_paragraph(line[2:], style="Heading 1")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string(colors['heading1'])
            
            continue
        
        # If no heading has been processed yet, skip content lines
        if not content_block_started:
            continue
        
        # Bold Text
        if re.match(r"\*\*(.*?)\*\*", line):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r"\*\*(.*?)\*\*", r"\1", line))
            run.bold = True
            
        # Bulleted Lists
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            
        # Numbered Lists
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(line, style="List Number")
            
        # Data for graphs (simple format: "GRAPH: label1,value1;label2,value2")
        elif line.startswith("GRAPH:"):
            data_str = line[6:].strip()
            items = data_str.split(";")
            labels = []
            values = []
            
            for item in items:
                parts = item.split(",")
                if len(parts) == 2:
                    labels.append(parts[0].strip())
                    values.append(float(parts[1].strip()))
            
            if labels and values:
                data_for_graphs.append((labels, values))
                
        # Normal Text
        else:
            doc.add_paragraph(line)
    
    # Generate and add graphs if data is available
    for i, (labels, values) in enumerate(data_for_graphs):
        plt.figure(figsize=(8, 4))
        plt.bar(labels, values, color=f'#{colors["accent"]}')
        plt.title(f"Data Visualization {i+1}")
        plt.tight_layout()
        
        # Save the graph
        graph_filename = f"temp_graph_{i}.png"
        plt.savefig(graph_filename)
        plt.close()
        
        # Add to document
        doc.add_picture(graph_filename, width=Inches(6))
        graph_paragraph = doc.paragraphs[-1]
        graph_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add page borders with user-specified color
    for section in doc.sections:
        # Add page borders - all sides
        section_properties = section._sectPr
        page_borders = OxmlElement('w:pgBorders')
        page_borders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # single line
            border.set(qn('w:sz'), '4')       # thickness in 1/8 points (4 = 0.5pt)
            border.set(qn('w:space'), '24')   # space in points
            border.set(qn('w:color'), colors['border'])  # Use user-specified border color
            page_borders.append(border)
            section_properties.append(page_borders)

    # Add header
    for section in doc.sections:
        header = section.header

        # Create a table in the header with one row and two columns
        table = header.add_table(rows=3, cols=2, width=Inches(6))
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set the width of the columns
        table.columns[0].width = Inches(1.5)  # Adjust as needed for the logo
        table.columns[1].width = Inches(4.5)  # Adjust as needed for the text

        # Add the logo to the first cell
        cell_logo = table.cell(0, 0)
        cell_logo_para = cell_logo.paragraphs[0]
        cell_logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run_logo = cell_logo_para.add_run()
        run_logo.add_picture("temp_logo.png", width=Pt(50))  # Smaller logo size

        # Add the text to the second cell
        cell_text = table.cell(0,1)
        cell_text_para = cell_text.paragraphs[0]
        cell_text_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run_text = cell_text_para.add_run("Proposal from Consultadd Public Services")
        run_text.font.underline = True
    
    # Add page numbers in footer with user-specified color
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"{company_name} Case Study | Page "
        
        # Style the footer text with user's footer color
        for run in footer_para.runs:
            run.font.color.rgb = RGBColor.from_string(colors['footer'])
        
        # Add page number field
        run = footer_para.add_run()
        run.font.color.rgb = RGBColor.from_string(colors['footer'])
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        run = footer_para.add_run(" of ")
        run.font.color.rgb = RGBColor.from_string(colors['footer'])
        
        run = footer_para.add_run()
        run.font.color.rgb = RGBColor.from_string(colors['footer'])
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "NUMPAGES"
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        run = footer_para.add_run(" | Confidential")
        run.font.color.rgb = RGBColor.from_string(colors['footer'])
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save document
    doc.save(docx_file)
    print(f"✅ DOCX saved as {docx_file}")
    
    # Convert to PDF
    pdf_file = output_file
    try:
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf", docx_file, "--outdir", os.path.dirname(pdf_file)], check=True)
        print(f"✅ PDF saved as {pdf_file}")
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")
    
    # Cleanup temporary files
    try:
        if logo_path and os.path.exists(logo_path):
            os.remove(logo_path)
            print("✅ Removed temporary logo file")
    except Exception as e:
        print(f"⚠️ Failed to clean up temporary logo file: {e}")